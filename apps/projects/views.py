from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from docx import Document
from .forms import ProjectForm
from .models import Project
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_project_doc(request, project_id):
    project = get_object_or_404(Project, id=project_id)
    team = project.team_assigned

    document = Document()

    title = document.add_heading(level=1)
    title_run = title.add_run(f'Proyecto: {project.title}')
    title_run.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar el título

    def add_bold_paragraph(doc, label, value):
        """Función para agregar párrafos con etiquetas en negrita."""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(label)
        run.bold = True
        paragraph.add_run(f' {value}')

    add_bold_paragraph(document, 'Manager:', project.tutor)
    add_bold_paragraph(document, 'Objetivo:', project.goal)
    add_bold_paragraph(document, 'Descripción:', project.description)
    add_bold_paragraph(document, 'Inicio:', project.start_date or "Sin definir")
    add_bold_paragraph(document, 'Entrega:', project.end_date or "Sin definir")
    add_bold_paragraph(document, 'Estado:', "Activo" if project.is_active else "Inactivo")
    add_bold_paragraph(document, 'Fase:', project.get_project_phase_display())

    # **Enlaces del Proyecto**
    if project.project_url or project.project_tiktok or project.project_instagram:
        document.add_heading('Enlaces del Proyecto', level=2)
        if project.project_url:
            add_bold_paragraph(document, 'Página Web:', project.project_url)
        if project.project_tiktok:
            add_bold_paragraph(document, 'TikTok:', project.project_tiktok)
        if project.project_instagram:
            add_bold_paragraph(document, 'Instagram:', project.project_instagram)

    document.add_heading('Integrantes del Equipo', level=2)
    if team and team.members.exists():
        for member in team.members.all():
            document.add_paragraph(f'{member.first_name} {member.last_name} | {member.career}', style='ListBullet')
    else:
        document.add_paragraph("No hay equipo asignado.")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=proyecto_{project.title}.docx'
    document.save(response)
    return response

@login_required
def create_project(request):
    team = request.user.team_assigned

    if not team or team.created_by != request.user:
        return redirect('dashboard')

    if team.project_assigned:
        return redirect('dashboard')

    if request.method == 'POST':
        form = ProjectForm(request.POST)
        if form.is_valid():
            project = form.save(user=request.user)  # Guardamos correctamente

            team.project_assigned = project  # Asignamos el proyecto al equipo
            team.save()  # Guardamos el equipo

            return redirect('dashboard')
    else:
        form = ProjectForm()

    return render(request, 'projects/create_project.html', {'form': form})

from django.shortcuts import render, get_object_or_404
from .models import Project

@login_required
def project_detail(request, project_id):
    project = get_object_or_404(Project, id=project_id)

    return render(request, 'projects/project_detail.html', {'project': project})


@login_required
def project_edit(request, project_id):
    project = get_object_or_404(Project, id=project_id)

    if project.created_by != request.user:
        return redirect('project_detail', project_id=project.id)

    if request.method == 'POST':
        form = ProjectForm(request.POST, instance=project)
        if form.is_valid():
            form.save(request.user)
            return redirect('project_detail', project_id=project.id)
    else:
        form = ProjectForm(instance=project)

    return render(request, 'projects/project_edit.html', {'form': form, 'project': project})


# API de verificacion
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json
from .services import find_similar_projects

@csrf_exempt
def check_project_similarity(request):
    """Verifica si un proyecto tiene similitud con otro en la BD."""
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            title = data.get("title", "")
            goal = data.get("goal", "")
            description = data.get("description", "")

            if not title or not description:
                return JsonResponse({"error": "Se requieren título y descripción."}, status=400)

            similar_projects = find_similar_projects(title, goal, description, threshold=0.8)

            return JsonResponse({"similar_projects": similar_projects})
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"error": "Método no permitido"}, status=405)